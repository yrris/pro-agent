"""附件处理（M8）：上传附件 → 多模态消息构造与入模型前展开。

设计要点（docs/08 §4）：
- **checkpoint 防膨胀**：state 里的 HumanMessage 只存自定义 `pro_attachment` 引用块
  （resource_key + 元信息），绝不存 base64；think 节点入模型前经 `expand_attachment_blocks`
  第三道投影（repair → history 裁剪 → expand）按需下载展开。
- **全量替换不变量**：expand 之后消息里不得残留任何 pro_attachment 块——provider
  对未知块类型直接 400（Anthropic），且残留会随 checkpoint 污染后续轮。
- **vision 门控**：仅 vision provider（anthropic）把图片展开为 base64 image 块；
  DeepSeek 等纯文本模型降级为文本占位。图片尺寸闸 4.5MB（Claude 单图 ~5MB 上限，
  与上传 20MB 上限不是一回事）。
- 下载走认知面自持的 MinIO 凭据（M4 起已有），带受限 LRU（小对象才缓存）。
"""

from __future__ import annotations

import base64
import io
import logging
import threading
from collections import OrderedDict
from typing import Any, Callable, Iterable, Optional, Sequence

from langchain_core.messages import AnyMessage, HumanMessage

from cognition.config import Settings, get_settings

logger = logging.getLogger(__name__)

# 自定义引用块类型（只存在于 checkpoint/state，永不直达 provider）。
ATTACHMENT_BLOCK_TYPE = "pro_attachment"

# 支持图像理解的 provider（能力表；M9 扩 provider 时在此登记）。
VISION_PROVIDERS = {"anthropic"}

# 展开期单图字节上限（Claude 单图 ~5MB，留余量）。
MAX_IMAGE_BYTES = int(4.5 * 1024 * 1024)

# 下载 LRU：最多 8 项、单对象 ≤6MB 才缓存（防大图吃内存）。
_CACHE_MAX_ITEMS = 8
_CACHE_MAX_BYTES = 6 * 1024 * 1024


def supports_vision(provider: str) -> bool:
    """该 provider 是否支持图像理解（executor 角色 resolved provider）。"""
    return (provider or "").lower() in VISION_PROVIDERS


def is_image(mime: str) -> bool:
    return (mime or "").lower().startswith("image/")


def normalize_attachments(attachments: Iterable[Any]) -> list[dict[str, Any]]:
    """proto Attachment / dict → 统一 dict（后续全链路只认这个形状）。"""
    out: list[dict[str, Any]] = []
    for a in attachments or []:
        if isinstance(a, dict):
            rk = str(a.get("resource_key", "") or "")
            fn = str(a.get("file_name", "") or "")
            mt = str(a.get("mime_type", "") or "")
            size = int(a.get("size", 0) or 0)
        else:
            rk = str(getattr(a, "resource_key", "") or "")
            fn = str(getattr(a, "file_name", "") or "")
            mt = str(getattr(a, "mime_type", "") or "")
            size = int(getattr(a, "size", 0) or 0)
        if rk:
            out.append({"resource_key": rk, "file_name": fn or rk.rsplit("/", 1)[-1], "mime_type": mt, "size": size})
    return out


def _human_size(n: int) -> str:
    if n < 1024:
        return f"{n}B"
    if n < 1024 * 1024:
        return f"{n / 1024:.0f}KB"
    return f"{n / 1024 / 1024:.1f}MB"


def attachment_note(attachments: Sequence[dict], ingested_names: Sequence[str] = ()) -> str:
    """附件清单注记（进消息文本，模型据此知道有什么、去哪找）。"""
    if not attachments:
        return ""
    items = "、".join(f"{a['file_name']}（{a['mime_type'] or '未知类型'}, {_human_size(a['size'])}）" for a in attachments)
    note = f"〔用户上传附件：{items}〕"
    if ingested_names:
        note += f"\n〔其中 {'、'.join(ingested_names)} 的内容已存入你的知识库，可用 knowledge_search 工具检索引用〕"
    return note


def build_attachment_message(
    query: str, attachments: Sequence[dict], ingested_names: Sequence[str] = ()
) -> HumanMessage:
    """构造带附件的用户消息：文本块（query+清单注记）+ 图片的 pro_attachment 引用块。

    非图片附件不产块（内容经知识库/引用块路径供给）；无图片时退化为纯文本 content
    （不走块路径，减少不必要的多模态分支）。
    """
    text = query if not attachments else f"{query}\n\n{attachment_note(attachments, ingested_names)}"
    image_atts = [a for a in attachments if is_image(a["mime_type"])]
    if not image_atts:
        return HumanMessage(content=text)
    blocks: list[Any] = [{"type": "text", "text": text}]
    for a in image_atts:
        blocks.append(
            {
                "type": ATTACHMENT_BLOCK_TYPE,
                "resource_key": a["resource_key"],
                "file_name": a["file_name"],
                "mime_type": a["mime_type"],
                "size": a["size"],
            }
        )
    return HumanMessage(content=blocks)


class MinioDownloader:
    """按 resource_key 读对象字节（惰性 import + 受限 LRU）。失败抛异常由调用方降级。"""

    def __init__(self, settings: Optional[Settings] = None) -> None:
        self._settings = settings or get_settings()
        self._client: Any = None
        self._cache: OrderedDict[str, bytes] = OrderedDict()

    def _get_client(self):
        if self._client is None:
            from minio import Minio  # 惰性：离线单测不需要该依赖

            s = self._settings
            self._client = Minio(
                s.minio_endpoint, access_key=s.minio_access_key,
                secret_key=s.minio_secret_key, secure=s.minio_secure,
            )
        return self._client

    def __call__(self, resource_key: str) -> bytes:
        cached = self._cache.get(resource_key)
        if cached is not None:
            self._cache.move_to_end(resource_key)
            return cached
        resp = self._get_client().get_object(self._settings.minio_bucket, resource_key)
        try:
            data = resp.read()
        finally:
            resp.close()
            resp.release_conn()
        if len(data) <= _CACHE_MAX_BYTES:
            self._cache[resource_key] = data
            while len(self._cache) > _CACHE_MAX_ITEMS:
                self._cache.popitem(last=False)
        return data


# —— 上传附件 → 知识库（run 前同步入库，read-your-writes）——

# 文本类判定（进知识库的内容源）：mime 前缀/精确匹配 + 扩展名兜底（浏览器对 md/csv
# 的 mime 报告不稳定）。
_TEXT_MIME_EXACT = {"application/json", "application/x-ndjson", "application/xml"}
_TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".json", ".log", ".xml", ".yaml", ".yml"}
PDF_MIME = "application/pdf"

# 单文件入库文本上限（约 200k 字符）：防超大文档拖慢 run 启动；超限截断并注记。
MAX_INGEST_CHARS = 200_000

# —— C 批（docs/13）：扫描版 PDF 逐页 OCR ——
# 扫描页判定阈值：页文本 strip 后不足该字符数即判为扫描页。不用纯判空——扫描件常被
# pypdf 提出页眉/页码几个字符的碎渣；30 字符以下的"文本页"即便误判去 OCR，转写结果
# 也不会更差。模块级常量便于 monkeypatch/调优（docs/13 §3.1）。
SCAN_TEXT_THRESHOLD = 30

# 每份 PDF 的 OCR **尝试**页数上限（vision 调用次数硬上限）：每扫描页一次 vision 调用
# 且发生在 run 前同步预步，超大扫描件靠此护栏保住首轮响应。按"尝试"而非"成功"计数
#（评审#5/#18）：转写失败/返回空的页同样消耗配额——否则 key 限流/空白图册场景下每页
# 都白付一次栅格化+vision 往返，护栏恰在最需要止损时失效。超限在文末如实注记而非静默
# 截断（docs/13 §2/§3.5）。
MAX_OCR_PAGES = 20

# OCR 采用阈值：转写结果 strip 后不足该字符数视为"未采用"——保留原页文本、不置
# ocr_used（但已消耗尝试配额）。防止文本 PDF 的短页（扉页/章节隔页，<30 字符被判为
# 扫描页）OCR 出空/碎渣后误触发 ocr_used=True，把整份文档的幂等寻址从内容哈希切到
# dedup_seed——既有已入库文本 PDF 重传会全量漂出新 point（评审#6，docs/13 §3.3 红线）。
MIN_OCR_TEXT_CHARS = 10


# —— D1 批（docs/15）：PDF 文本页结构化提取层（表格 / 多栏 / 公式）参数 ——
# 仅作用于"文本页"（strip 后 >= SCAN_TEXT_THRESHOLD，原走原样保留分支）；纯 prose 单栏
# 无表格页保持 plain 逐字节原样（回归红线 test_pure_text_pdf_output_identical，docs/15
# §3.2）。每个子步骤惰性 import pdfplumber、逐步 try/except 降级回退 plain，绝不炸 run。

# 双栏检测：单侧词数下限——太少不判栏（避免把稀疏页/短标题页误判为两栏）。
MIN_COLUMN_WORDS = 6
# 双栏 gutter 中带宽度（占页宽比例）：中线两侧各半带；带内被任何词横占即判"非对半两栏"
#（贯穿全宽的单行 prose 会横占中带 → 判否 → 保 plain，守住红线）。
COLUMN_GUTTER_BAND = 0.06
# 表格采纳下限：低于 2 行 / 2 列的"退化表"忽略——只认真正有行列结构的表。评审 #9 起
# 抽表仅用 lines 策略（需绘制线），本护栏作为 defense-in-depth 兜底过滤病态单行/单列表。
TABLE_MIN_ROWS = 2
TABLE_MIN_COLS = 2
# 公式启发式：整页数学符号（数学运算符 / 希腊字母 / 上下标等 Unicode 段）计数达此阈值才
# 触发整页 vision 兜底。保守取值，避免中文（CJK 段不计入）/普通 prose 误触发。
FORMULA_MIN_SYMBOLS = 8


def is_text_like(mime: str, file_name: str = "") -> bool:
    m = (mime or "").lower()
    if m.startswith("text/") or m in _TEXT_MIME_EXACT:
        return True
    name = (file_name or "").lower()
    return any(name.endswith(ext) for ext in _TEXT_EXTS)


_MAX_OFFICE_UNCOMPRESSED = 100 * 1024 * 1024  # office(zip) 解压总大小上限，防 zip-bomb DoS


def _office_zip_safe(data: bytes) -> bool:
    """docx/xlsx 本质是 zip：解析前校验解压总大小（20MB 压缩包可膨胀到 GB 级）。"""
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            total = sum(i.file_size for i in zf.infolist())
        return total <= _MAX_OFFICE_UNCOMPRESSED
    except Exception:  # noqa: BLE001 — 非法 zip 交给下游解析报错
        return True


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def is_docx(mime: str, file_name: str = "") -> bool:
    return mime == _DOCX_MIME or file_name.lower().endswith(".docx")


def is_xlsx(mime: str, file_name: str = "") -> bool:
    return mime == _XLSX_MIME or file_name.lower().endswith(".xlsx")


def is_pdf(mime: str, file_name: str = "") -> bool:
    return (mime or "").lower() == PDF_MIME or (file_name or "").lower().endswith(".pdf")


# pypdfium2/PDFium 官方声明 "inherently not thread-safe"（PDFium 内部全局状态；
# ctypes 外调释放 GIL 即真并发）：Run 前置入库与 IngestDocument 均经 asyncio.to_thread
# 跑在默认线程池的不同 worker 上，两个线程同时进入 pdfium C 调用是未定义行为——典型
# 后果是内存损坏/段错误，这类原生崩溃不是 Python 异常，except 兜不住，整个认知面进程
# 连同所有在飞 run 一起死掉（评审#4）。模块级互斥锁把全部 pdfium 调用串行化（官方推荐
# 的多线程用法）；栅格化毫秒级、慢路径 vision OCR 在锁外，串行化对吞吐影响可忽略。
_PDFIUM_LOCK = threading.Lock()

# 渲染前像素预算（约 A4@300DPI 再上浮）：PDF MediaBox 允许 14400×14400pt，scale=2 下
# 单页位图可达 GB 级，且缓冲在 pdfium C 层渲染时直接分配——MAX_IMAGE_BYTES 字节闸只
# 作用于编码后的 PNG/JPEG，拦不住渲染缓冲本身（评审#7/#17）。必须在 render 之前按页
# 点尺寸反解安全 scale，对齐 _office_zip_safe 的"解析前防护"先例。
MAX_RENDER_PIXELS = 12_000_000


def _rasterize_pdf_page(data: bytes, page_index: int, *, scale: float = 2.0) -> Optional[bytes]:
    """把 PDF 单页栅格化为位图字节（C 批 docs/13 §3.2）。

    惰性 import pypdfium2（自带 pdfium 二进制 wheel，python-slim 镜像零系统依赖；
    对齐本仓库解析依赖"函数内 import + 异常降级"的既定风格）。
    渲染前预算闸（评审#7/#17）：先用 page.get_size()（点）反解安全
    scale = min(scale, sqrt(MAX_RENDER_PIXELS/(w*h)))，让位图分配上限有界；夹到
    0.5 以下仍超预算（异常巨幅 MediaBox）→ 跳过该页返回 None（调用方保留原文本），
    绝不让 GB 级缓冲先于任何检查被提交。
    尺寸闸：PNG 超 MAX_IMAGE_BYTES（转写器拒转线）→ 降 scale 重渲一次 → 仍超转
    JPEG q85（有损但转写足够）→ 还超只能放弃——否则高 DPI 页会被转写器静默拒转。
    调用方按魔数嗅探 PNG/JPEG 决定 mime。任何失败降级返回 None，不炸调用方。
    线程安全（评审#4）：pdfium 全局状态非线程安全，从 PdfDocument 打开到 close 的
    全部调用持模块级 _PDFIUM_LOCK 串行化（见常量注释）。
    """
    try:
        import math

        import pypdfium2 as pdfium  # 惰性：未装/坏字节都降级

        def _encode(im: Any, fmt: str, **kw: Any) -> bytes:
            out = io.BytesIO()
            im.save(out, format=fmt, **kw)
            return out.getvalue()

        with _PDFIUM_LOCK:
            pdf = pdfium.PdfDocument(data)
            try:
                page = pdf[page_index]
                try:
                    # 渲染前像素预算：按页点尺寸夹紧 scale（分配发生在 render 内的 C 层，
                    # 这里是唯一能在分配之前挡住它的位置）。
                    w_pt, h_pt = page.get_size()
                    area = max(float(w_pt) * float(h_pt), 1.0)
                    scale = min(scale, math.sqrt(MAX_RENDER_PIXELS / area))
                    if scale < 0.5:  # 0.5 兜底 ≈ 36DPI，再低无转写价值——巨幅页直接跳过
                        logger.warning(
                            "pdf page %s too large to rasterize (%.0fx%.0f pt), skipped",
                            page_index + 1, w_pt, h_pt,
                        )
                        return None
                    img = page.render(scale=scale).to_pil()
                    png = _encode(img, "PNG")
                    if len(png) <= MAX_IMAGE_BYTES:
                        return png
                    # 高 DPI 页超闸：降半 scale 重渲（0.5 兜底；scale 已过预算闸 ≥0.5，
                    # 降半后仍 ≤ 预算上限，不会重新超预算）。
                    img = page.render(scale=max(scale / 2, 0.5)).to_pil()
                    png = _encode(img, "PNG")
                    if len(png) <= MAX_IMAGE_BYTES:
                        return png
                    jpg = _encode(img.convert("RGB"), "JPEG", quality=85)
                    return jpg if len(jpg) <= MAX_IMAGE_BYTES else None
                finally:
                    page.close()
            finally:
                pdf.close()
    except Exception as exc:  # noqa: BLE001 — 栅格化失败降级，调用方保留该页原文本
        logger.warning("pdf page rasterize failed (page %s): %s", page_index + 1, exc)
        return None


# —— D1 批（docs/15 §3.1）：PDF 文本页结构化提取辅助 ——
# 全部子步骤对 None/异常一律降级回退 plain（红线：结构不识别时逐字节保留原文本）。


def _open_plumber(data: bytes) -> Optional[Any]:
    """惰性打开 pdfplumber PDF（docs/15 §3.3）：未装/坏字节 → None，文本页回退 plain。

    模块级函数便于离线单测 monkeypatch 注入 fake plumber（伪造 extract_tables /
    extract_words / crop，不触真依赖）。
    """
    try:
        import pdfplumber  # 惰性：未装则整层降级，对齐仓库解析依赖既定风格

        return pdfplumber.open(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 — 打开失败降级，文本页回退 plain 原样
        logger.warning("pdfplumber open failed: %s", exc)
        return None


def _plumber_page(plumber: Any, idx: int) -> Optional[Any]:
    """安全取 pdfplumber 第 idx 页；plumber 为 None / 越界 / 异常 → None（降级）。"""
    if plumber is None:
        return None
    try:
        pages = plumber.pages
        return pages[idx] if 0 <= idx < len(pages) else None
    except Exception as exc:  # noqa: BLE001 — 取页失败降级回退 plain
        logger.warning("pdfplumber page %s access failed: %s", idx + 1, exc)
        return None


def _table_to_markdown(table: Sequence[Sequence[Any]]) -> Optional[str]:
    """二维单元格数组 → markdown 表格（继承 docx `" | "` 先例，对 split_text 透明）。

    退化表（去空行后 < TABLE_MIN_ROWS 行或 < TABLE_MIN_COLS 列）返回 None——作为 lines
    策略抽表的 defense-in-depth 兜底，过滤病态单行/单列表。单元格内换行/竖线做转义避免破坏表结构。
    """
    def _cell(v: Any) -> str:
        s = "" if v is None else str(v)
        return s.replace("\r", " ").replace("\n", " ").replace("|", r"\|").strip()

    rows = [[_cell(c) for c in (row or [])] for row in table if row is not None]
    rows = [r for r in rows if any(r)]  # 去空行
    if len(rows) < TABLE_MIN_ROWS:
        return None
    ncol = max(len(r) for r in rows)
    if ncol < TABLE_MIN_COLS:
        return None
    rows = [r + [""] * (ncol - len(r)) for r in rows]  # 补齐短行到等列
    lines = ["| " + " | ".join(rows[0]) + " |",
             "| " + " | ".join(["---"] * ncol) + " |"]
    lines += ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join(lines)


def _extract_tables_md(ppage: Any) -> Optional[list[str]]:
    """pdfplumber 抽表 → markdown 串列表（**仅默认 lines/explicit 策略：需真实绘制的表格线**）；
    无有效表 / 异常 → None（回退 plain）。退化表由 _table_to_markdown 兜底过滤。

    评审 #9 红线修复：**去掉原 text 策略兜底**。text 策略靠词的 x 对齐推断行列，会把对齐的
    多行普通 prose（每行几个词、行间 y 对齐）偶合识别成 >=2 行 >=2 列的伪表——退化表护栏
    （TABLE_MIN_ROWS/COLS=2）拦不住它；命中后整页被追加一段乱切 markdown 表 → 输出不再逐字节
    一致、build_ingestor 据 transformed 置 dedup_seed，直接破坏 docs/13 §3.3 + docs/15 §3.2
    『纯文本页逐字节零回归』红线（既有已入库文档漂出新 point、旧点残留）。真正的表格有绘制线
    （默认 lines 策略即可检出），prose 无线——故只认有线的表；对齐的多栏/prose 不再被误判，
    正确落到下游『多栏 gutter 检测』或『保持 plain 原样』分支，红线不破而真表仍被抽取。
    """
    try:
        # 仅默认 lines 策略：需真实划线、误报率极低；不再用 text 策略兜底（评审 #9 红线）。
        tables = ppage.extract_tables() or []
        mds = [md for t in tables if (md := _table_to_markdown(t))]
        return mds or None
    except Exception as exc:  # noqa: BLE001 — 抽表失败降级回退 plain
        logger.warning("pdfplumber extract_tables failed: %s", exc)
        return None


def _two_column_text(ppage: Any) -> Optional[str]:
    """两栏对半启发式（docs/15 §3.1）：检出页中部贯通垂直空白 gutter → 裁左右子页各
    extract_text → 先左后右拼接还原阅读顺序；非两栏 / 异常 → None（回退 plain）。

    gutter 判定：以页宽中线为中心取一条窄带，若任何词横占该带即判"非对半两栏"——贯穿
    全宽的单行 prose 会横占中带，据此被排除（守红线）；再要求左右两侧各有足量词，排除
    单栏左/右对齐页。最小版只处理对半两栏（覆盖多数报告/论文正文区，docs/15 §2）。
    """
    try:
        words = ppage.extract_words() or []
        if len(words) < MIN_COLUMN_WORDS * 2:
            return None
        width = float(ppage.width)
        height = float(ppage.height)
        mid = width / 2.0
        band = width * COLUMN_GUTTER_BAND
        lo, hi = mid - band / 2, mid + band / 2
        left = right = 0
        for w in words:
            x0, x1 = float(w["x0"]), float(w["x1"])
            if x0 < hi and x1 > lo:
                return None  # 词横占 gutter 中带 → 非对半两栏（含全宽单行 prose）
            if (x0 + x1) / 2 < mid:
                left += 1
            else:
                right += 1
        if left < MIN_COLUMN_WORDS or right < MIN_COLUMN_WORDS:
            return None  # 一侧太空 → 单栏左/右对齐，非两栏
        lt = (ppage.crop((0, 0, mid, height)).extract_text() or "").strip()
        rt = (ppage.crop((mid, 0, width, height)).extract_text() or "").strip()
        merged = f"{lt}\n{rt}".strip()
        return merged or None
    except Exception as exc:  # noqa: BLE001 — 裁列失败降级回退 plain
        logger.warning("pdfplumber two-column detect failed: %s", exc)
        return None


def _is_math_char(c: str) -> bool:
    """字符是否落在"数学符号"Unicode 段（公式启发式用）。CJK 段不计入 → 中文不误触发。"""
    o = ord(c)
    return (
        0x2200 <= o <= 0x22FF  # Mathematical Operators（∑∫∏√∞≈≠≤≥∂∇∈∉⊂⊃∪∩…）
        or 0x2A00 <= o <= 0x2AFF  # Supplemental Mathematical Operators
        or 0x27C0 <= o <= 0x27EF  # Misc Mathematical Symbols-A
        or 0x2980 <= o <= 0x29FF  # Misc Mathematical Symbols-B
        or 0x2070 <= o <= 0x209F  # Superscripts and Subscripts（上标/下标）
        or 0x0370 <= o <= 0x03FF  # Greek（α β γ δ θ λ μ π σ φ ψ ω…公式常用）
        or c in "±×÷"  # 落在 Latin-1 补充段的常见数学符号
    )


def _looks_formula_heavy(text: str) -> bool:
    """整页数学符号计数达 FORMULA_MIN_SYMBOLS → 判为公式/复杂版式页（docs/15 §3.1）。"""
    return sum(1 for c in text if _is_math_char(c)) >= FORMULA_MIN_SYMBOLS


def _extract_pdf_text(
    data: bytes,
    file_name: str = "",
    transcribe: Optional[Callable[[bytes, str], Optional[str]]] = None,
) -> tuple[Optional[str], bool]:
    """PDF 逐页提取 + 扫描页 OCR（C 批 docs/13）+ 文本页结构化提取（D1 docs/15）。
    返回 (文本, 页文本是否经过非确定性变换)。

    逐页判定：
    - **文本页**（strip 后 >= SCAN_TEXT_THRESHOLD）走 D1（docs/15 §3.1）结构化提取层：
      ① 表格（pdfplumber extract_tables，**仅 lines 策略：需真实绘制线**，评审 #9 去 text
      兜底）非空 → plain 正文 + 各表 markdown（〔第N页·表格〕标记）；② 两栏对半 → 裁左右列
      先左后右重排；③ 命中数学符号
      密度启发式 → 整页 vision 兜底（复用扫描 OCR 链路与配额）；④ 纯 prose 单栏无表格 →
      **plain 逐字节原样**（回归红线 test_pure_text_pdf_output_identical，docs/15 §3.2）。
      pdfplumber 未装/任一步异常一律降级回退 plain，绝不炸 run。
    - **扫描页**（< 阈值）且有转写器 → 栅格化 → vision OCR → 转写 strip 后 >=
      MIN_OCR_TEXT_CHARS 才采用，替换为 `〔第N页·OCR〕\\n转写文本`；碎渣结果视为未采用
      （评审#6：防短文本页误触发切换幂等寻址）。无转写器时整体行为与现状一致。
    OCR **尝试**页数达 MAX_OCR_PAGES 后不再 OCR（扫描页 + 公式兜底页共享该硬上限；送
    transcribe 即计、无论成败——vision 调用次数硬上限，评审#5/#18），文末追加超限注记。
    返回的第二个 bool（D1 docs/15 §3.2 泛化）：只要有页被**非确定性变换**过（扫描/公式
    OCR 或表格/多栏结构化重排）即 True——供 build_ingestor 决定 dedup_seed：变换后文本
    与旧 plain 逐字节不同（OCR 更天然漂移），必须切到文件字节寻址幂等（docs/13 §3.3）。
    """
    plumber: Any = None
    plumber_opened = False
    try:
        from pypdf import PdfReader  # 惰性：未装/损坏 pdf 都降级

        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        ocr_attempted = 0  # 已尝试 OCR 的页数：调用 transcribe 即计、无论成败（配额计数）
        scan_total = 0  # 检出的扫描页总数（仅在有转写器时统计，用于超限注记）
        limit_skipped = 0  # 因尝试上限未 OCR 的扫描页数（>0 才追加注记）
        transformed = False  # D1（docs/15 §3.2）：是否有页经非确定性变换（OCR / 结构化重排）
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            # —— 文本页：D1 结构化提取层（表格 / 多栏 / 公式），无结构则 plain 逐字节原样 ——
            if len(text.strip()) >= SCAN_TEXT_THRESHOLD:
                if not plumber_opened:  # 惰性打开：扫描件/无文本页不付 pdfplumber 解析成本
                    plumber_opened = True
                    plumber = _open_plumber(data)
                ppage = _plumber_page(plumber, idx)
                # ① 表格：非空 → plain 正文 + 各表 markdown（〔第N页·表格〕标记同 OCR 款）
                mds = _extract_tables_md(ppage) if ppage is not None else None
                if mds:
                    parts = [text] + [f"〔第{idx + 1}页·表格〕\n{md}" for md in mds]
                    pages.append("\n".join(parts))
                    transformed = True
                    continue
                # ② 两栏对半 → 裁左右子页先左后右重排阅读顺序
                two_col = _two_column_text(ppage) if ppage is not None else None
                if two_col is not None:
                    pages.append(two_col)
                    transformed = True
                    continue
                # ③ 公式：命中数学符号密度启发式 → 整页 vision 兜底（复用扫描 OCR 配额上限）
                if (transcribe is not None and _looks_formula_heavy(text)
                        and ocr_attempted < MAX_OCR_PAGES):
                    fx: Optional[str] = None
                    try:
                        img = _rasterize_pdf_page(data, idx)
                        if img is not None:
                            img_mime = "image/png" if img.startswith(b"\x89PNG") else "image/jpeg"
                            ocr_attempted += 1  # 尝试即计（与扫描页同款硬上限语义）
                            fx = transcribe(img, img_mime)
                    except Exception as exc:  # noqa: BLE001 — 兜底失败降级保留 plain，不炸整份
                        logger.warning("pdf formula-page OCR failed for %s (page %s): %s",
                                       file_name, idx + 1, exc)
                    if fx and len(fx.strip()) >= MIN_OCR_TEXT_CHARS:
                        pages.append(f"〔第{idx + 1}页·OCR〕\n{fx}")
                        transformed = True
                        continue
                # ④ 纯 prose 单栏无表格：plain 逐字节原样（回归红线，docs/15 §3.2）
                pages.append(text)
                continue
            # —— 扫描页（< 阈值）——
            if transcribe is None:
                pages.append(text)  # 无转写器：整体行为与现状一致
                continue
            scan_total += 1
            if ocr_attempted >= MAX_OCR_PAGES:
                limit_skipped += 1
                pages.append(text)  # 超上限：保留原文本（可能为空）
                continue
            ocr: Optional[str] = None
            try:
                img = _rasterize_pdf_page(data, idx)
                if img is not None:
                    # 栅格化按尺寸可能降级为 JPEG，按魔数嗅探 mime（docs/13 §3.2）。
                    img_mime = "image/png" if img.startswith(b"\x89PNG") else "image/jpeg"
                    # 尝试即消耗配额（在调用前递增：transcribe 抛异常同样已付出一次
                    # vision 往返）——失败页不计数会让护栏在 key 限流/空白图册等最需要
                    # 止损的场景下失效，vision 调用次数实际无上限（评审#5/#18）。
                    ocr_attempted += 1
                    ocr = transcribe(img, img_mime)
            except Exception as exc:  # noqa: BLE001 — 单页 OCR 失败降级保留原文本，不拖垮整份
                logger.warning("pdf page OCR failed for %s (page %s): %s", file_name, idx + 1, exc)
            # 采用闸（评审#6）：碎渣转写（strip 后 < MIN_OCR_TEXT_CHARS，如扉页/隔页
            # OCR 出空或回显几个字符）不采用——保留原页文本、不计入 transformed，文本
            # PDF 的内容寻址幂等不被切到 seed；尝试配额已在上面扣除。
            if ocr and len(ocr.strip()) >= MIN_OCR_TEXT_CHARS:
                transformed = True
                pages.append(f"〔第{idx + 1}页·OCR〕\n{ocr}")
            else:
                pages.append(text)
        merged = "\n".join(pages).strip()
        if limit_skipped:
            # 注记语义（评审#5）：X=检出的扫描页总数，N=已尝试 OCR 的页数（含失败页）。
            note = f"〔注：扫描页共{scan_total}页，仅前{ocr_attempted}页已尝试OCR〕"
            merged = f"{merged}\n{note}" if merged else note
        return (merged or None), transformed
    except Exception as exc:  # noqa: BLE001 — pdf 解析失败降级跳过（坏 PDF 必须仍返回 None）
        logger.warning("pdf text extract failed for %s: %s", file_name, exc)
        return None, False
    finally:
        if plumber is not None:  # D1：释放 pdfplumber 打开的底层文件句柄
            try:
                plumber.close()
            except Exception:  # noqa: BLE001 — 关闭失败无害
                pass


def extract_text(
    data: bytes,
    mime: str,
    file_name: str = "",
    *,
    transcribe: Optional[Callable[[bytes, str], Optional[str]]] = None,
) -> Optional[str]:
    """从附件字节提取纯文本；不可提取/失败返回 None（调用方跳过，不炸 run）。

    transcribe：图片 OCR 转写器（B.2）——给了才对 image/* 转文本，否则图片仍返回 None；
    C 批（docs/13）起 PDF 分支同样复用它做扫描页逐页 OCR（内部走 _extract_pdf_text，
    丢弃 ocr_used 信号；需要该信号的调用方用 extract_text_ex）。
    """
    if is_image(mime) and transcribe is not None:
        return transcribe(data, mime)
    if is_pdf(mime, file_name):
        return _extract_pdf_text(data, file_name, transcribe)[0]
    if is_docx(mime, file_name):
        if not _office_zip_safe(data):
            logger.warning("docx 解压过大（疑似 zip-bomb），跳过: %s", file_name)
            return None
        try:
            from docx import Document  # 惰性：python-docx

            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:  # 表格逐行拼为管道分隔文本
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            text = "\n".join(parts).strip()
            return text or None
        except Exception as exc:  # noqa: BLE001 — 解析失败降级跳过
            logger.warning("docx text extract failed for %s: %s", file_name, exc)
            return None
    if is_xlsx(mime, file_name):
        if not _office_zip_safe(data):
            logger.warning("xlsx 解压过大（疑似 zip-bomb），跳过: %s", file_name)
            return None
        try:
            from openpyxl import load_workbook  # 惰性：openpyxl

            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts: list[str] = []
            acc = 0  # 累积字符上限：read_only 流式，边读边停，不等全表读完
            for ws in wb.worksheets:
                parts.append(f"# 工作表: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v) for v in row]
                    if any(c.strip() for c in cells):
                        line = ",".join(cells)
                        parts.append(line)
                        acc += len(line)
                        if acc > MAX_INGEST_CHARS:
                            parts.append("…（超长截断）")
                            wb.close()
                            return "\n".join(parts).strip()
            wb.close()
            text = "\n".join(parts).strip()
            return text or None
        except Exception as exc:  # noqa: BLE001
            logger.warning("xlsx text extract failed for %s: %s", file_name, exc)
            return None
    if is_text_like(mime, file_name):
        try:
            return data.decode("utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            return None
    return None


def extract_text_ex(
    data: bytes,
    mime: str,
    file_name: str = "",
    *,
    transcribe: Optional[Callable[[bytes, str], Optional[str]]] = None,
) -> tuple[Optional[str], bool]:
    """extract_text 的扩展版（C 批 docs/13 §3.3）：额外返回"页文本是否被非确定性变换过"。

    该 bool（D1 docs/15 §3.2 泛化）对 PDF 扫描/公式 OCR **或**表格/多栏结构化重排置
    True——build_ingestor 据此把幂等键从内容寻址切到 dedup_seed=md5(文件字节)（变换后
    文本与旧 plain 逐字节不同、OCR 更天然漂移，防重复入库）。图片分支恒 False：图片的
    种子由 build_ingestor 既有的 ocr_image 门负责，不动既有逻辑。
    头部两个分支的判定顺序必须与 extract_text 保持一致（image 先于 pdf）。
    """
    if is_image(mime) and transcribe is not None:
        return transcribe(data, mime), False
    if is_pdf(mime, file_name):
        return _extract_pdf_text(data, file_name, transcribe)
    return extract_text(data, mime, file_name, transcribe=transcribe), False


def build_image_transcriber(settings: Settings) -> Optional[Callable[[bytes, str], Optional[str]]]:
    """构建图片 OCR/转写器（B.2）：用 anthropic vision 把图片转成可入库文本。

    **不依赖 executor 角色**（默认 deepseek 无视觉）——只要有 anthropic key 就可用；
    无 key 返回 None（调用方据此不放行图片入库）。返回的闭包同步阻塞（在 to_thread 内跑）。
    """
    if not getattr(settings, "anthropic_api_key", None):
        return None
    from langchain_core.messages import HumanMessage

    from cognition.providers.anthropic_provider import build_anthropic_chat

    model = build_anthropic_chat(settings, max_tokens=2000)

    def _transcribe(data: bytes, mime: str) -> Optional[str]:
        if len(data) > MAX_IMAGE_BYTES:
            return None  # 过大图不 OCR（成本/超限）
        import base64 as _b64

        b64 = _b64.b64encode(data).decode("ascii")
        msg = HumanMessage(content=[
            {"type": "text", "text": "请把这张图片里的所有文字/表格/图表信息完整转写成纯文本"
                                      "（保留结构，用于检索）。只输出转写内容，不要额外说明。"},
            {"type": "image", "source_type": "base64", "data": b64, "mime_type": mime or "image/png"},
        ])
        try:
            out = model.invoke([msg])
            text = out.content if isinstance(out.content, str) else str(out.content)
            return text.strip() or None
        except Exception as exc:  # noqa: BLE001 — OCR 失败降级跳过，不拖垮入库
            logger.warning("image OCR failed: %s", exc)
            return None

    return _transcribe


def build_ingestor(
    settings: Optional[Settings] = None,
    *,
    downloader: Optional[Callable[[str], bytes]] = None,
    store: Any = None,
    embedder: Any = None,
    sparse: Any = None,
    transcribe: Optional[Callable[[bytes, str], Optional[str]]] = None,
) -> Callable[[list[dict], str], list[str]]:
    """构建附件入库器（装配期一次；I/O 依赖可注入供离线测试）。

    返回同步可调用 `(att_dicts, kb_id) -> 入库文件名列表`——servicer 在 Run 里经
    `asyncio.to_thread` 调用（embedder/下载同步阻塞，绝不能在 grpc.aio 事件循环上裸跑）。
    幂等：`ingest(stable_ids=True)` 内容寻址，同文件重传/重试不重复入库。
    """
    settings = settings or get_settings()
    if store is None or embedder is None or sparse is None:
        from cognition.rag.factory import build_embedder, build_sparse, build_store

        store = store or build_store(settings)
        embedder = embedder or build_embedder(settings)
        sparse = sparse or build_sparse(settings)
    dl = downloader or MinioDownloader(settings)
    # B.2：图片 OCR 转写器（有 anthropic key 才启用；显式传 transcribe 覆盖，供离线测试注入 fake）。
    _transcribe = transcribe if transcribe is not None else build_image_transcriber(settings)

    def _ingest_attachments(attachments: list[dict], kb_id: str) -> list[str]:
        from cognition.rag.ingest import ingest

        if not kb_id:
            return []  # kb_id 空=无隔离全库，宁可不入
        docs: list[dict] = []
        names: list[str] = []
        for a in attachments or []:
            mime, fname = a.get("mime_type", ""), a.get("file_name", "")
            # 图片仅在有 OCR 转写器时放行入库（B.2）；否则仍走多模态/占位路径。
            ocr_image = _transcribe is not None and is_image(mime)
            if not (is_text_like(mime, fname) or is_pdf(mime, fname) or is_docx(mime, fname)
                    or is_xlsx(mime, fname) or ocr_image):
                continue
            try:
                data = dl(a["resource_key"])
            except Exception as exc:  # noqa: BLE001 — 单文件失败不拖垮其余
                logger.warning("attachment download failed for ingest %s: %s", a.get("resource_key"), exc)
                continue
            # C 批（docs/13）：改调 _ex 版拿"非确定性变换"信号（D1 docs/15 §3.2 泛化：
            # 扫描/公式 OCR 或表格/多栏结构化重排均置 True）。
            text, ocr_used = extract_text_ex(data, mime, fname, transcribe=_transcribe)
            if not text or not text.strip():
                continue
            if len(text) > MAX_INGEST_CHARS:
                text = text[:MAX_INGEST_CHARS] + "\n…（超长截断）"
            doc = {"text": text, "file_name": fname, "source_id": a["resource_key"]}
            if ocr_image or ocr_used:
                # 变换后文本与旧 plain 逐字节不同（图片整图 OCR=ocr_image 既有门；PDF 扫描/
                # 公式 OCR、表格/多栏重排=ocr_used，C 批 docs/13 §3.3 + D1 docs/15 §3.2）——
                # 用文件字节内容哈希做稳定幂等种子，保证同文件重传/重跑原地 upsert，不因文本
                # 漂移/重排而重复入库（评审#8）。纯 prose 单栏无表格 PDF ocr_used=False 绝不
                # 带 seed：保持内容寻址幂等不变（回归红线）。
                import hashlib as _hl

                doc["dedup_seed"] = _hl.md5(data).hexdigest()
            docs.append(doc)
            names.append(fname)
        if docs:
            ingest(docs, kb_id, store=store, embedder=embedder, sparse=sparse, stable_ids=True)
        return names

    return _ingest_attachments


def _placeholder_block(file_name: str, reason: str) -> dict:
    return {"type": "text", "text": f"[图片附件 {file_name}（{reason}，未注入图像内容）]"}


def expand_attachment_blocks(
    messages: list[AnyMessage],
    *,
    downloader: Callable[[str], bytes],
    vision: bool,
    max_image_bytes: int = MAX_IMAGE_BYTES,
) -> list[AnyMessage]:
    """把消息里的 pro_attachment 引用块展开为真实内容（入模型前最后一道投影）。

    vision provider：下载 → base64 image 块（LangChain 标准块，ChatAnthropic 原生消费）；
    非 vision / 下载失败 / 超尺寸：降级文本占位。**保证输出不残留任何 pro_attachment 块**。
    只读投影：不回写 state/checkpoint。
    """
    out: list[AnyMessage] = []
    for m in messages:
        content = getattr(m, "content", None)
        if not isinstance(content, list) or not any(
            isinstance(b, dict) and b.get("type") == ATTACHMENT_BLOCK_TYPE for b in content
        ):
            out.append(m)
            continue
        new_blocks: list[Any] = []
        for b in content:
            if not (isinstance(b, dict) and b.get("type") == ATTACHMENT_BLOCK_TYPE):
                new_blocks.append(b)
                continue
            fname = str(b.get("file_name", "附件"))
            mime = str(b.get("mime_type", ""))
            key = str(b.get("resource_key", ""))
            if not vision or not is_image(mime):
                new_blocks.append(_placeholder_block(fname, "当前模型不支持图像理解"))
                continue
            try:
                data = downloader(key)
            except Exception as exc:  # noqa: BLE001 — 读取失败降级占位，不炸 run
                logger.warning("attachment download failed for %s: %s", key, exc)
                new_blocks.append(_placeholder_block(fname, "图片读取失败"))
                continue
            if len(data) > max_image_bytes:
                new_blocks.append(_placeholder_block(fname, f"图片超过 {_human_size(max_image_bytes)} 模型上限"))
                continue
            new_blocks.append(
                {
                    "type": "image",
                    "source_type": "base64",
                    "data": base64.b64encode(data).decode("ascii"),
                    "mime_type": mime,
                }
            )
        out.append(m.model_copy(update={"content": new_blocks}))
    return out
